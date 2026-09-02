import os, sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders_xml = f"""
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
    </w:tblBorders>
    """
    tblPr.append(parse_xml(borders_xml))

def style_table(table, col_widths, headers, data, header_bg="1E3A8A", alt_bg="F8FAFC"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Header Row
    hdr_row = table.rows[0]
    trPr = hdr_row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:tblHeader'))
    
    for i, h_text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = col_widths[i]
        set_cell_background(cell, header_bg)
        set_cell_margins(cell, top=120, bottom=120, left=130, right=130)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h_text)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    # Data Rows
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        bg_color = alt_bg if (r_idx % 2 == 1) else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=90, bottom=90, left=130, right=130)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(str(val))
            run.font.name = '맑은 고딕'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(30, 41, 59)

def add_p(doc, text, bold_prefix=None, space_after=4, line_spacing=1.15):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if bold_prefix:
        r_b = p.add_run(bold_prefix)
        r_b.font.name = '맑은 고딕'
        r_b.font.bold = True
        r_b.font.size = Pt(9.5)
        r_b.font.color.rgb = RGBColor(15, 23, 42)
    if text:
        r = p.add_run(text)
        r.font.name = '맑은 고딕'
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_bullet(doc, title, desc):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r_b = p.add_run(f"{title}: ")
    r_b.font.name = '맑은 고딕'
    r_b.font.bold = True
    r_b.font.size = Pt(9)
    r_b.font.color.rgb = RGBColor(15, 23, 42)
    r = p.add_run(desc)
    r.font.name = '맑은 고딕'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(51, 65, 85)
    return p

def build_urs_docx(output_path):
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Title & Subtitle
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_title.add_run("User Requirements Specification (URS)\n")
    r_sub.font.name = '맑은 고딕'
    r_sub.font.size = Pt(14)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(37, 99, 235)
    
    r_main = p_title.add_run("Factory Operation AI Agent 사용자 요구사항 명세서")
    r_main.font.name = '맑은 고딕'
    r_main.font.size = Pt(20)
    r_main.font.bold = True
    r_main.font.color.rgb = RGBColor(15, 23, 42)
    p_title.paragraph_format.space_after = Pt(12)
    
    # Doc Info Table
    doc_info_data = [
        ["문서 번호", "URS-FOA-001"],
        ["문서명", "Factory Operation AI Agent 사용자 요구사항 명세서 (URS)"],
        ["버전", "v0.13"],
        ["작성자", "자동화AX솔루션팀"],
        ["검토자 / 승인자", "공장 운영기술팀 / 시스템 기획팀"],
        ["최종 개정일", "2026-08-21"],
        ["문서 상태", "Approved (배포용)"],
        ["적용 범위", "스마트 팩토리 제조·물류 자동화 공정 운영 및 보전 지원"]
    ]
    t_info = doc.add_table(rows=len(doc_info_data)+1, cols=2)
    style_table(t_info, [Inches(1.8), Inches(4.8)], ["항목", "내용"], doc_info_data, header_bg="0F172A")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # Revision History
    h1_rev = doc.add_heading("개정 이력 (Revision History)", level=1)
    h1_rev.paragraph_format.space_before = Pt(12)
    h1_rev.paragraph_format.space_after = Pt(4)
    
    rev_data = [
        ["v0.00", "2026-03-20", "초기 URS 초안 작성", "자동화AX"],
        ["v0.01", "2026-08-14", "스마트 팩토리 범용 도메인 용어 정제 및 모니터링 용어 변경", "자동화AX"],
        ["v0.02", "2026-08-18", "요구사항 ID 체계(FR, DR, IR, PR, SR, OR) 구조화", "자동화AX"],
        ["v0.03", "2026-08-20", "안전 원격 제어 및 정비 일지 삭제, 6대 핵심 가치 체계로 정제", "자동화AX"],
        ["v0.04", "2026-08-20", "실시간 모니터링 오해 표현 정제 및 AI Agent HOW 메커니즘 명시", "자동화AX"],
        ["v0.05", "2026-08-20", "태그(Tag) 관리 기능 제거 및 단순화", "자동화AX"],
        ["v0.06", "2026-08-21", "관제 대시보드 및 전문 뷰어 요소 전면 제거, Chat-First 정체성 확립", "자동화AX"],
        ["v0.07", "2026-08-21", "FN-CTL-02 삭제 및 FN-CTL-01을 FN-EQP-05로 편입, 핵심 가치 개편", "자동화AX"],
        ["v0.08", "2026-08-21", "FN-ADV-03(트러블슈팅 Tree) 삭제 및 운영 산출물 3대 표준 문서 정예화", "자동화AX"],
        ["v0.09", "2026-08-21", "개인화 워크스페이스를 플랫폼 기능으로 분리 및 4대 비즈니스 핵심 가치 확정", "자동화AX"],
        ["v0.10", "2026-08-21", "WIP 용어 전면 제거 및 공정 재고/재공품으로 일괄 표준화 반영", "자동화AX"],
        ["v0.11", "2026-08-21", "카테고리 관련 내용 전면 제거 및 워크스페이스 기능 정예화", "자동화AX"],
        ["v0.12", "2026-08-21", "Scoped RAG를 '프로젝트 전용 매뉴얼 맞춤 검색'으로 사용자 친화적 용어 변경", "자동화AX"],
        ["v0.13", "2026-08-21", "템플릿 관련 내용 전면 제거 및 기능/산출물 명세 정제", "자동화AX"]
    ]
    t_rev = doc.add_table(rows=len(rev_data)+1, cols=4)
    style_table(t_rev, [Inches(0.8), Inches(1.1), Inches(3.7), Inches(1.0)], ["버전", "개정일", "주요 개정 내용", "작성자"], rev_data, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 1. 개요
    doc.add_heading("1. 개요 (Introduction)", level=1)
    
    # 1.1 목적
    doc.add_heading("1.1 목적 (Purpose)", level=2)
    add_p(doc, "본 문서는 스마트 팩토리 제조·물류 공정 운영 및 보전 지원을 위한 인공지능 에이전트(가칭 Factory Operation AI Agent, 이하 '에이전트')의 사용자 요구사항을 정의한다. 본 요구사항 명세서는 이후 기능 명세(FS), 시스템 설계(DS), 및 테스트 검증(QA)의 공식 기준선으로 활용된다.")
    
    # 1.2 범위
    doc.add_heading("1.2 범위 (Scope)", level=2)
    add_p(doc, "본 프로젝트의 업무 및 시스템 구축 범위는 다음과 같다.")
    
    scope_items = [
        ("대상 업무", "물류 자동화 설비(AGV, RGV, 컨베이어, AS/RS, S/C, Gantry 등) 보전 SOP 및 정비 매뉴얼 질의응답, 설비 가동 현황 및 가동률/알람 통계 분석, 자재 및 공정 재고 체류시간 추적, AI 시스템 운영 지원 산출물(PM 체크리스트, 시운전 가이드, BCP 절차서) 자동 생성, 개인화 프로젝트 워크스페이스 기반 전용 매뉴얼 맞춤 검색 관리"),
        ("시스템 경계", "웹/채팅 인터페이스(Chat UI), Multi-SLM 코어 엔진(Clas·Plan·Proc), 지식 Vector DB(SOP·매뉴얼 임베딩), RDB(계정·이력) 및 외부 데이터 단일 통로인 중계 서버(Relay Server) REST API 클라이언트"),
        ("대상 사용자", "현장 물류·설비 보전 엔지니어, 자재/창고 및 공정 물류 관리자, 생산 기술/공장 운영 관리자, 시스템 관리자"),
        ("대상 사이트", "제조·물류 공장 자동화 구역 (원자재/부자재 창고, 공정 간 버퍼 구역, 완제품 적재 및 이송 대기 구역)")
    ]
    for title, desc in scope_items:
        add_bullet(doc, title, desc)
        
    add_p(doc, "\n[범위 제외 사항 (Out of Scope)]", bold_prefix="")
    out_scope_items = [
        "MES, WMS, SCADA, 설비 PLC 자체의 신규 개발 및 레거시 소스 수정",
        "에이전트에 의한 설비 PLC/DB 직접 물리적 제어 및 임의 SQL(Text-to-SQL) 생성",
        "설비 완전 무인 자동 운전 판단 및 공정 품질 예측/PdM(예지보전) 모델링"
    ]
    for item in out_scope_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"✕ {item}")
        p.paragraph_format.line_spacing = 1.15

    # 1.3 용어 및 약어
    doc.add_heading("1.3 용어 및 약어 정의 (Glossary)", level=2)
    terms_data = [
        ["URS", "User Requirements Specification — 사용자 요구사항 명세서"],
        ["Factory Operation AI Agent", "제조 및 공정 물류 현장의 보전 엔지니어와 관리자를 위한 대화형 스마트 팩토리 지능화 에이전트"],
        ["Multi-SLM", "단일 거대 모델 대신 의도 분류(Clas), 작업 계획(Plan), 실행/가공(Proc)으로 특화 분화된 모듈형 경량 언어 모델 엔진"],
        ["Clas-SLM", "사용자 질의 프롬프트의 의도 및 도메인을 신속 분류하는 전담 모델"],
        ["Plan-SLM", "분류된 의도에 따라 소작업 순서, 지식 검색 키워드 및 API 호출 계획을 수립하는 전담 모델"],
        ["Proc-SLM", "지식 검색 결과 요약, 수신 데이터의 통계 연산 및 대화형 표/카드 피드백을 가공하는 전담 모델"],
        ["RAG", "Retrieval-Augmented Generation — 외부 지식(SOP, PDF 매뉴얼)을 벡터 검색하여 환각 없이 정확한 근거 기반으로 답변하는 기술"],
        ["프로젝트 전용 지식 한정 검색", "전체 공장 DB 대신 특정 프로젝트 워크스페이스 내 문서로 검색 범위를 한정하여 타 설비 혼선을 차단하는 맞춤형 정밀 검색 기술"],
        ["중계 서버 (Relay Server)", "현장 설비 PLC/DB 직접 접근을 차단하고 규격화된 데이터 조회 REST API를 제공하는 단일 관문"],
        ["Function Calling", "자연어 질의에서 파라미터(기간, 호기 등)를 추출하여 중계 서버 REST API 규격에 맞게 자동 호출하는 SLM 기능"],
        ["공정 재고 (재공품)", "공정 간 버퍼 구역에 적재되어 후속 공정 투입을 대기 중인 중간 생산품 및 버퍼 재고"],
        ["체류 시간 (Dwell Time)", "공정 지연 방지 및 품질 관리를 위해 버퍼 구역에서 관리되는 공정 재고의 대기 시간"],
        ["FIFO", "First In First Out — 선입선출 자재 및 재고 관리 원칙"],
        ["SOP", "Standard Operating Procedure — 설비 알람/장애 발생 시 표준 점검 및 조치 수순 가이드"],
        ["Grounding", "응답 내용을 검색된 문서 원문 근거에 한정하고 원문 출처(페이지 딥링크)를 함께 제시하여 환각을 원천 억제하는 방식"],
        ["Project Workspace", "개인/팀별로 담당 설비 매뉴얼(PDF)과 도면을 사전 등록하여 세션 초월 영구 재참조하는 전용 지식 공간 (플랫폼 편의 기능)"],
        ["BCP", "Business Continuity Plan — 네트워크 단선 등 비상 상황 시 현장 수동 운전(Manual Override) 전환 절차서"],
        ["PM", "Preventive Maintenance — 설비 고장을 예방하기 위해 정기적(일/주/월)으로 수행하는 예방 보전 점검 체크리스트"]
    ]
    t_terms = doc.add_table(rows=len(terms_data)+1, cols=2)
    style_table(t_terms, [Inches(2.2), Inches(4.4)], ["용어 / 약어", "정의 및 설명"], terms_data, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 2. 시스템 개요 및 4대 핵심 가치
    doc.add_heading("2. 시스템 개요 및 4대 핵심 가치 (System Overview)", level=1)
    
    doc.add_heading("2.1 시스템 목표 및 4대 핵심 제공 가치", level=2)
    add_p(doc, "본 시스템은 상시 감시용 관제 화면이 아닌, 자연어 질의응답을 통해 제조·물류 설비 보전 지식과 공정 데이터를 즉시 제공하는 대화형 AX(AI Transformation) 에이전트로서 다음 4대 비즈니스 핵심 가치를 제공한다.")
    
    core_values_data = [
        ["1. 보전 지식 & SOP 안내", "FN-SOP-01~04", "설비 알람 발생 시 지식 검색으로 1~N단계 점검 수순, 부품 규격, 도면 위치를 대화형으로 단계별 즉시 안내 (MTTR 단축)"],
        ["2. 설비 성능 & 가동률/알람 분석", "FN-EQP-01~05", "자연어 질의로 설비별 가동 실적, 시간가동률(OEE), 병목 구간, 비가동 원인 및 알람 통계를 중계 서버 API로 집계하여 리포트 제공 (가동률 극대화)"],
        ["3. 자재 & 공정 재고 추적", "FN-INV-01~03", "자재 창고 적재율, 유효기간(D-day), FIFO 위반 경고 및 공정 버퍼 체류 시간 초과(지연) 항목을 대화형으로 추적 (자재 변질/정체 방지)"],
        ["4. 운영 지원 산출물 자동 제공", "FN-ADV-01~03", "표준 마스터 데이터, 정기 PM 체크리스트, 5단계 시운전 가이드, BCP 수동 전환 절차서 등 표준 운영 문서 자동 작성 (문서 공수 절감)"]
    ]
    t_cv = doc.add_table(rows=len(core_values_data)+1, cols=3)
    style_table(t_cv, [Inches(1.8), Inches(1.2), Inches(3.8)], ["번호 / 핵심 제공 가치", "세부 기능 ID", "주요 제공 가치 및 설명"], core_values_data, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 2.2 사용자 및 역할
    doc.add_heading("2.2 사용자 및 역할 정의 (User Roles & RBAC)", level=2)
    roles_data = [
        ["현장 물류/설비 보전 엔지니어", "User", "물류 설비 알람 대응 및 유지보수 점검을 수행하는 현장 엔지니어", "설비 상태 조회, SOP 및 정비 매뉴얼 질의, 멀티모달 PDF 업로드, 개인 워크스페이스 관리"],
        ["자재/창고 및 공정 물류 관리자", "User", "자재 적재 현황, 공정 버퍼 체류 시간, 공정 간 물류 흐름을 관리하는 담당자", "재고·체류 시간·공정 대기열 조회 및 분석, FIFO 위반 및 반출 지연 알림 확인"],
        ["공장 운영 및 생산 기술 관리자", "User", "설비 가동률 분석, 표준 운영 문서 및 점검 절차를 수립하는 관리자", "설비 가동률/비가동 분석 리포트 조회, PM 체크리스트/시운전 가이드/BCP 절차서 생성"],
        ["시스템 관리자 (Admin)", "Admin", "AI 에이전트 시스템 계정, 권한, 전사 지식 데이터를 운영·유지보수하는 담당자", "사용자 계정 및 권한 관리, 전사 지식(SOP/매뉴얼) 갱신, 연동 설정 관리, 시스템 모니터링"]
    ]
    t_roles = doc.add_table(rows=len(roles_data)+1, cols=4)
    style_table(t_roles, [Inches(1.8), Inches(0.8), Inches(1.9), Inches(2.1)], ["역할명", "RBAC", "주요 담당 업무", "주요 접근 권한"], roles_data, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 2.3 4대 핵심 가치별 세부 기능 명세
    doc.add_heading("2.3 4대 핵심 가치별 세부 기능 상세 명세 (Detailed Features)", level=2)
    
    # 2.3.1 가치 1
    doc.add_heading("2.3.1 [가치 1] 물류 설비 보전 지식 & SOP 매뉴얼 안내", level=3)
    sop_features = [
        ["FN-SOP-01", "설비 고장 조치 SOP 단계별 가이드", "• 물류 설비 알람 발생 시 1~N단계 현장 표준 점검 및 조치 수순(SOP)을 자연어로 단계별 안내\n• LOTO(안전 잠금) 및 전원 차단 안전 수칙 강조", "• 자연어 질의 (예: 'AGV 3호기 라인 이탈 알람(ERR-104) 조치 SOP')\n• 설비 모델명, 에러 코드", "1. Clas-SLM 의도 분류\n2. SOP 지식 DB 검색\n3. Proc-SLM 1~N단계 수순 요약", "• 1~N단계 체크포인트 카드\n• 안전 유의사항 배너\n• 참조 SOP 딥링크"],
        ["FN-SOP-02", "설비 정비 및 유지보수 매뉴얼 지원", "• 정기 점검 주기, 소모품/부품 교체 가이드, 체결 토크 기준, 분해/조립 수순 질의응답\n• 순정 부품 품번(Part Number) 안내", "• 부품명, 설비 모델명 질의 (예: 'S/C 와이어로프 권장 교체 주기 및 규격')", "1. 정비 매뉴얼 청크 검색\n2. 토크/주기/품번 정형화 추출", "• 부품 정비 스펙 표(토크, 규격, 주기)\n• 분해/조립 순서도"],
        ["FN-SOP-03", "고장 조치 이력 탐색 (Troubleshooting)", "• 과거 동종 설비/알람 발생 사례 및 보전 엔지니어의 조치 이력 검색을 통한 원인 규명", "• 알람 코드, 이상 현상 텍스트 (예: 'RGV 감속 이상 빈발 원인 및 조치 사례')", "1. 과거 조치 이력 검색\n2. 원인별 빈도 및 성공사례 랭킹", "• 유사 장애 Top 3 조치 사례\n• 원인별 점유율(%) 표"],
        ["FN-SOP-04", "멀티모달 매뉴얼 지식 융합", "• 사용자가 업로드한 PDF, Office, 도면 파일(최대 50MB)을 실시간 파싱·임베딩하여 문서 내 세부 조치법 검색", "• 정비 매뉴얼(PDF), 도면 파일 첨부 + 'E-203 조치법 찾아줘'", "1. PDF 텍스트/도표 OCR 파싱\n2. 지식 임베딩 매칭\n3. 원본 좌표 및 텍스트 요약", "• 에러 전용 조치 요약문\n• 원문 페이지 딥링크\n• 배선도/도면 안내"]
    ]
    t_sop = doc.add_table(rows=len(sop_features)+1, cols=6)
    style_table(t_sop, [Inches(0.9), Inches(1.1), Inches(1.4), Inches(1.1), Inches(1.0), Inches(1.1)], ["ID", "기능명", "기능 정의", "입력", "작동 메커니즘 (HOW)", "출력"], sop_features, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 2.3.2 가치 2
    doc.add_heading("2.3.2 [가치 2] 물류 설비 성능 및 가동률/알람 분석", level=3)
    eqp_features = [
        ["FN-EQP-01", "이송 설비 이상 진단 및 현황 질의", "• AGV 및 무인 지게차의 운전 모드, 저배터리, 경로 이탈 및 장애 정지 등 이상 상태 질의응답", "• 질의 프롬프트 (예: 'AGV 가동 상태 및 배터리 20% 미만 차량 조회')", "1. 설비 그룹 및 조건 파싱\n2. 중계 서버 REST API 호출\n3. JSON 조건 필터링 및 서식 가공", "• 이송 설비 상태 표\n• 저배터리/이상 경고 하이라이트"],
        ["FN-EQP-02", "이송 시스템 처리 성능 및 병목 진단", "• Sorting C/V, Lifter, Diverter, RGV, S/C 라인별 시간당 처리량(TPH), 사이클 타임 비교 및 병목 구간 식별", "• 질의 프롬프트 (예: 'AS/RS 1~4호기 S/C 처리량 편차 및 속도 저하 분석')", "1. 실적 데이터 API 수집\n2. TPH 및 사이클 타임 편차 계산\n3. 병목 구간 자동 식별", "• 호기별 TPH/사이클 타임 비교 표\n• 병목/속도 저하 진단 코멘트"],
        ["FN-EQP-03", "적재/투입 설비 가동 상태 & 큐 정체 분석", "• 로더/언로더 가동 상태, 버퍼 대기 큐 정체 및 투입 지연 요인 분석", "• 질의 프롬프트 (예: '투입 로더 3번 가동 상태 및 잔여 투입 대기량')", "1. PLC 접점값 및 대기 큐 수집\n2. 설비 모드 및 대기 큐 과부하 진단", "• 설비 상태 요약 카드 (상태, 시간당 실적, 정체 진단)"],
        ["FN-EQP-04", "설비 가동률 및 비가동 원인 분석", "• 일/주/월간 설비별 시간가동률(OEE) 집계 및 비가동 원인(고장, 대기 등) 파레토 분석 리포트 생성", "• 질의 프롬프트 (예: '이번 주 RGV 라인 비가동 원인 Top 3 및 가동률')", "1. 기간/설비군 파라미터 추출\n2. 가동률/알람 로그 API 연동\n3. 비가동 요인 파레토 통계 분석", "• 가동률 집계 표\n• 비가동 원인 파레토 순위 카드\n• AI 원인 요약 및 개선 권고"],
        ["FN-EQP-05", "설비 알람 통계 및 빈발 장애 분석", "• 설비/호기/알람코드별 발생 빈도 및 정지 시간 통계 분석을 통한 만성 반복 알람 패턴 식별", "• 질의 프롬프트 (예: '최근 24시간 S/C 1호기 발생 알람 통계 및 빈발 에러')", "1. 알람 발생/해제 로그 API 수집\n2. 정지 시간 및 최다 빈출 코드 집계\n3. 반복성 알람 패턴 브리핑 생성", "• 알람 통계 요약 표\n• 빈발 알람 분석 브리핑 카드"]
    ]
    t_eqp = doc.add_table(rows=len(eqp_features)+1, cols=6)
    style_table(t_eqp, [Inches(0.9), Inches(1.1), Inches(1.4), Inches(1.1), Inches(1.0), Inches(1.1)], ["ID", "기능명", "기능 정의", "입력", "작동 메커니즘 (HOW)", "출력"], eqp_features, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 2.3.3 가치 3
    doc.add_heading("2.3.3 [가치 3] 자재 및 공정 재고 이상 추적", level=3)
    inv_features = [
        ["FN-INV-01", "자재 유효기간 및 FIFO 위반 추적", "• 자재 창고 구역별 선입선출(FIFO) 위반 검출 및 유효기간(D-day) 임박 자재 이상 추적·경고", "• 질의 프롬프트 (예: '원자재 2창고 유효기간 3일 이내 및 FIFO 위반 목록')", "1. WMS 재고 Lot 데이터 수신\n2. FIFO 위반 및 D-day 계산\n3. 위험도 순 정렬 및 경고 포맷팅", "• 자재 Lot별 현황 표 (입고일, D-day, 랙 번호)\n• FIFO 위반 경고 배너"],
        ["FN-INV-02", "공정 버퍼 정체 진단 & 체류 시간 분석", "• 공정 간 버퍼 재고 체류 시간 분석 및 기준 시간(예: 45분) 초과 정체 품목 식별·원인 진단", "• 질의 프롬프트 (예: '버퍼 구역에서 45분 초과된 공정 재고 목록과 원인')", "1. MES 버퍼 적재 품목/시각 수신\n2. 체류 시간 차이 연산 (45분 초과)\n3. 빈 캐리어/팔레트 동시 집계", "• 공정 재고 체류 분석 표\n• 주의/지연 인디케이터\n• 공 용기 잔여량 및 정체 진단"],
        ["FN-INV-03", "공정 간 물류 이송 지연 탐지 & 도착 예측", "• 공정 간 이동 중인 화물의 이송 지연 병목 구간 탐지 및 도착 예상 시각 산출", "• 질의 프롬프트 (예: '가공 1라인에서 조립 라인 이송 물류 현황 및 지연 여부')", "1. 이송 지시 및 운송 설비 위치 추적\n2. 표준 시간 대비 지연 구간 탐지\n3. 도착 예상 시각 산출", "• 이송 진단 카드 (화물ID, 운송AGV, 예상도착, 지연 경고)"]
    ]
    t_inv = doc.add_table(rows=len(inv_features)+1, cols=6)
    style_table(t_inv, [Inches(0.9), Inches(1.1), Inches(1.4), Inches(1.1), Inches(1.0), Inches(1.1)], ["ID", "기능명", "기능 정의", "입력", "작동 메커니즘 (HOW)", "출력"], inv_features, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 2.3.4 가치 4
    doc.add_heading("2.3.4 [가치 4] AI 시스템 운영 지원 산출물 자동 제공", level=3)
    adv_features = [
        ["FN-ADV-01", "표준 데이터 & 정기 PM 체크리스트 생성", "• 신규/개조 설비 도입 시 표준 데이터 매핑 추천 및 일/주/월간 정기 예방 보전(PM) 체크리스트 자동 작성", "• 산출물 요청 (예: '신규 도입된 RGV 5호기 주간 PM 체크리스트 작성해줘')", "1. 유사 모델 정비 매뉴얼 지식 검색\n2. 점검 부위/방법/기준 표준 표 가공", "• PM 체크리스트 양식 표 (점검 부위, 방법, 기준, 주기)\n• 표준 마스터 데이터 제안"],
        ["FN-ADV-02", "테스트 / 시운전 가이드 지원", "• 시스템 패치 또는 개조 시 단위/통합 테스트 케이스 및 5단계 시운전 절차서 초안 자동 수립", "• 산출물 요청 (예: '컨베이어 PLC 로직 수정 후 시운전 테스트 절차서')", "1. 수정 로직 및 인터록 영향 분석\n2. 5단계 시운전 수순 절차서 생성", "• 5단계 시운전 가이드 문서\n• 테스트 시나리오/사전조건/기대결과"],
        ["FN-ADV-03", "비상 대응(BCP) & 수동 전환 가이드", "• 상위 시스템 다운 또는 통신 장애 시 현장 제어반 수동 운전(Manual Override) 전환 수순 및 비상 연락망 안내", "• 산출물 요청 (예: 'WMS 서버 다운 시 창고 수동 입출고 BCP 절차')", "1. BCP SOP 표준 절차서 탐색\n2. 수동 운전 3단계 수순 및 연락망 출력", "• BCP 비상 절차서 (수동 전환 수순, 안전 확인 체크리스트, 핫라인)"]
    ]
    t_adv = doc.add_table(rows=len(adv_features)+1, cols=6)
    style_table(t_adv, [Inches(0.9), Inches(1.1), Inches(1.4), Inches(1.1), Inches(1.0), Inches(1.1)], ["ID", "기능명", "기능 정의", "입력", "작동 메커니즘 (HOW)", "출력"], adv_features, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 2.4 플랫폼 기반 편의 기능
    doc.add_heading("2.4 [플랫폼 기반 편의 기능] 개인화 프로젝트 워크스페이스 (Platform Enabler)", level=2)
    pws_features = [
        ["FN-PWS-01", "개인 / 팀별 프로젝트 워크스페이스", "• 개인/작업반 단위로 프로젝트 공간을 생성하고 자주 참조하는 매뉴얼(PDF), 도면을 영구 보관하여 세션 초월 재참조 지원", "• 프로젝트 생성 명령, 매뉴얼/도면 파일 업로드", "1. 워크스페이스 식별자(ID) 부여\n2. 전용 Vector DB 파싱 및 인덱싱", "• 프로젝트 지식 저장소 관리 패널\n• 업로드 문서 인덱스 목록"],
        ["FN-PWS-02", "프로젝트 전용 매뉴얼 맞춤 검색", "• 전체 공장 DB 대신 현재 활성화된 프로젝트 워크스페이스 내 문서로만 검색 범위를 한정하여 답변 정확도 극대화", "• 워크스페이스 활성화 토글 + 질의 프롬프트", "1. 지정 네임스페이스로 검색 범위 한정\n2. 타 설비 매뉴얼 혼선 차단 맞춤형 정밀 답변", "• 프로젝트 한정 맞춤형 정밀 답변\n• 전용 문서 출처 인덱스"]
    ]
    t_pws = doc.add_table(rows=len(pws_features)+1, cols=6)
    style_table(t_pws, [Inches(0.9), Inches(1.1), Inches(1.4), Inches(1.1), Inches(1.0), Inches(1.1)], ["ID", "기능명", "기능 정의", "입력", "작동 메커니즘 (HOW)", "출력"], pws_features, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 2.5 에이전트 운영 원칙 및 제약사항
    doc.add_heading("2.5 에이전트 운영 원칙 및 제약사항 (Operational Principles & Constraints)", level=2)
    principles = [
        ("Chat-First AI Agent 본질 및 뷰어/대시보드 배제 원칙", "본 시스템은 상시 감시용 SCADA 대시보드나 전문 그래픽 뷰어가 아니며, 대화형 인터페이스(Chat UI)를 통해 질의응답, 원인 진단, 통계 분석, 운영 산출물 생성을 지원하는 지능형 AX 에이전트임"),
        ("데이터 직접 접근 차단 원칙", "에이전트는 현장 제어반(PLC)이나 원시 DB에 직접 접근하는 것을 엄격히 금지하며, 발주사와 약속된 단일 통로인 중계 서버(Relay Server) REST API를 통해서만 안전하게 데이터를 조회함"),
        ("TAG (Text-to-SQL) 위험성 차단 원칙", "SLM이 임의의 DB Query(TAG)를 직접 생성하는 방식을 엄격히 금지하여 복잡한 레거시 DB 오작동 및 과부하 위험을 원천 차단함"),
        ("환각 억제 및 그라운딩(Grounding) 원칙", "응답 내용은 지식 검색된 원문 근거에 한정하며 원출처(문서명, 페이지)를 명시하고, 신뢰도 점수(0.70) 미만 시 추정 답변을 자동 차단함"),
        ("집계 보고 및 통계 표출 방식 제약", "가동률, 알람 통계 등은 별도 고정 대시보드가 아닌 대화창 내 인라인 동적 표(Dynamic Table) 및 요약 카드로 렌더링함")
    ]
    for title, desc in principles:
        add_bullet(doc, title, desc)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 3. 주요 사용자 시나리오
    doc.add_heading("3. 주요 사용자 시나리오 및 운영 흐름 (User Scenarios)", level=1)
    
    doc.add_heading("3.1 메인 유저 저니 (Core User Journey)", level=2)
    journeys = [
        ("설비 알람 탐지 ➔ SOP 탐색 ➔ 현장 점검 수순 가이드", "AGV 3호기 라인 이탈 알람 발생(ERR-104) ➔ 사용자가 자연어로 SOP 질의 ➔ Clas-SLM 및 지식 검색 기반으로 1~2단계 점검 수순(센서 이물질 제거 ➔ 제로점 재설정) 및 도면 링크 즉시 피드백 ➔ 현장 조치 진행"),
        ("멀티모달 매뉴얼(PDF) 파일 기반 고장 원인 및 정비 수순 탐색", "RGV 정비 매뉴얼 PDF(45MB) 첨부 + 'E-203 조치법 요약' 프롬프트 입력 ➔ 실시간 PDF 파싱 및 지식 분석 ➔ 전압 측정 3단계 수순 가이드 및 PDF 45페이지 딥링크 제공"),
        ("공정 버퍼 체류 지연 탐지 ➔ 대화형 목록 확인 및 출고 조치", "물류 관리자가 '버퍼 구역 45분 초과 공정 재고 목록' 질의 ➔ 중계 서버 연동 후 체류 시간 45분 초과 항목 탐지 ➔ 적재 랙 번호, 잔여 유효시간, 지연 경고 인디케이터 표출"),
        ("시스템 변경 및 신규 설비 도입 시 운영 산출물 자동 생성", "운영 관리자가 '신규 RGV 5호기 주간 PM 체크리스트 및 시운전 가이드' 요청 ➔ Proc-SLM이 표준 데이터 및 매뉴얼 지식 탐색 ➔ 표준 PM 점검표 및 5단계 시운전 절차서 자동 작성"),
        ("개인화 프로젝트 워크스페이스 등록 및 매뉴얼 연속 질의", "엔지니어가 'AGV 3호기 전용 공간' 프로젝트 생성 및 매뉴얼 업로드 ➔ 세션 초월 영구 보관 지식 기반으로 문서 재첨부 없이 즉시 맞춤형 정밀 질의응답")
    ]
    for title, desc in journeys:
        add_bullet(doc, title, desc)

    doc.add_heading("3.2 예외 및 시스템 안전 시나리오", level=2)
    exceptions = [
        ("프롬프트 필수 인자 누락 시 대화형 재확인", "사용자가 '알람 조치법 알려줘' 입력 (설비명 누락) ➔ 에이전트가 에러 대신 '어느 설비의 알람을 조회하시겠습니까? [1] AGV 3호기 [2] S/C 1호기' 선택 칩 UI 출력 ➔ 클릭 후 정상 탐색"),
        ("중계 서버 통신 장애 시 사용자 안내", "중계 서버 REST API 응답 지연/타임아웃(3초) 발생 시 오류 사실 안내 및 수동 관제실 확인 가이드 표출"),
        ("지식 검색 근거 부족 및 유사도 미달 시 답변 불가 안내 (환각 차단)", "질의에 대한 지식 검색 유사도가 0.70 미달 시 임의 생성 차단 및 '관련 보전 SOP 근거를 찾을 수 없습니다' 안내문 표출")
    ]
    for title, desc in exceptions:
        add_bullet(doc, title, desc)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 4. 프롬프트 및 SLM 처리 체계
    doc.add_heading("4. 프롬프트 및 Multi-SLM 처리 체계 (Engine Architecture)", level=1)
    
    doc.add_heading("4.1 프롬프트 라이프사이클 (Prompt Lifecycle)", level=2)
    add_p(doc, "사용자 질의는 [① 생성(Create) ➔ ② 식별(Identify: Clas-SLM) ➔ ③ 계획(Plan: Plan-SLM) ➔ ④ 실행(Execute: Proc-SLM & 지식검색/중계서버) ➔ ⑤ 피드백 표출 및 보관(Finalize)]의 5단계 파이프라인으로 안전하고 정확하게 처리된다.")
    
    doc.add_heading("4.2 Multi-SLM 역할 분담 체계", level=2)
    slm_roles = [
        ("Clas-SLM (Classification)", "사용자 질의 프롬프트의 의도 및 도메인을 신속 분류하는 전담 모델"),
        ("Plan-SLM (Planning)", "분류된 의도에 따라 소작업(Task) 순서, 지식 검색 키워드 및 중계 서버 REST API 호출 계획을 수립하는 전담 모델"),
        ("Proc-SLM (Processing)", "계획에 따라 지식 검색 결과 요약, 수신 데이터의 통계 연산 및 대화형 표/카드 피드백을 가공하는 전담 모델")
    ]
    for title, desc in slm_roles:
        add_bullet(doc, title, desc)
        
    doc.add_heading("4.3 중계 서버 연동 및 인터페이스 확장성", level=2)
    add_p(doc, "에이전트는 사전 검증된 중계 서버 REST API 인터페이스 범위 내에서만 도구 호출(Function Calling)을 수행하여 안전성을 극대화하며, 신규 설비나 API 추가 시 모델 재학습 없이 '인터페이스 레지스트리(API Schema)' 설정 파일 갱신만으로 즉시 연동을 확장할 수 있는 동적 구조를 채택한다.")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 5. 상세 요구사항 명세
    doc.add_heading("5. 상세 요구사항 명세 (Requirements Specification)", level=1)
    
    # 5.1 기능 요구사항 FR (22개)
    doc.add_heading("5.1 기능 요구사항 (Functional Requirements - FR)", level=2)
    fr_data = [
        ["FR-001", "시스템은 사용자가 자연어 채팅 인터페이스(텍스트/음성/멀티모달 파일)를 통해 질의하고 피드백을 확인할 수 있어야 한다.", "M", "Concept 1.6, 7.1"],
        ["FR-002", "시스템은 Multi-SLM(Clas-Plan-Proc) 엔진을 통해 의도 분류, 실행 계획 수립, 지식 검색 및 중계 서버 연동 처리를 수행해야 한다.", "M", "Concept 3.4"],
        ["FR-003", "시스템은 물류 설비(AGV, RGV, S/C, 컨베이어 등) 알람 발생 시 1~N단계 현장 표준 점검 및 조치 수순(SOP)을 단계별 안내해야 한다.", "M", "FN-SOP-01"],
        ["FR-004", "시스템은 정기 점검 주기, 소모품/부품 교체 가이드, 토크 기준, 분해/조립 수순 및 순정 부품 품번 질의응답을 제공해야 한다.", "M", "FN-SOP-02"],
        ["FR-005", "시스템은 과거 동종 설비/알람 발생 사례 및 보전 엔지니어의 실제 조치 이력 검색을 통한 트러블슈팅 원인 규명을 지원해야 한다.", "M", "FN-SOP-03"],
        ["FR-006", "시스템은 사용자가 업로드한 정비 매뉴얼(PDF/도면)을 실시간 파싱·임베딩하여 문서 내 세부 조치법 검색 및 딥링크를 제공해야 한다.", "M", "FN-SOP-04"],
        ["FR-007", "시스템은 물류 이송 설비(AGV/무인지게차)의 운전 모드, 배터리 잔량, 이상 상태를 대화형으로 질의·진단할 수 있어야 한다.", "M", "FN-EQP-01"],
        ["FR-008", "시스템은 Sorting C/V, Lifter, Diverter, RGV, S/C 라인별 시간당 처리량(TPH), 사이클 타임 비교 및 병목 구간을 식별·진단해야 한다.", "M", "FN-EQP-02"],
        ["FR-009", "시스템은 투입 로더/언로더 작동 상태, 투입 대기량 및 버퍼 대기 큐 정체 요인을 분석할 수 있어야 한다.", "M", "FN-EQP-03"],
        ["FR-010", "시스템은 일간/주간/월간 설비별 시간가동률(OEE) 집계 및 비가동 원인 파레토 분석 리포트를 대화창 내 동적 표로 렌더링해야 한다.", "M", "FN-EQP-04"],
        ["FR-011", "시스템은 자재 창고 구역별 적재율, 유효기간(D-day) 카운트다운 관리 및 선입선출(FIFO) 위반 경고를 제공해야 한다.", "M", "FN-INV-01"],
        ["FR-012", "시스템은 공정 간 버퍼 적재량, 체류 시간 분석(정상/주의/지연 인디케이터) 및 공 팔레트/용기 잔여량을 제공해야 한다.", "M", "FN-INV-02"],
        ["FR-013", "시스템은 공정 간 이송 중인 화물의 이동 트래킹 및 이송 지연 병목 구간을 탐지해야 한다.", "M", "FN-INV-03"],
        ["FR-014", "시스템은 설비/호기/알람코드별 발생 빈도 및 정지 시간 통계 분석을 통해 빈발 알람 요약 브리핑을 제공해야 한다.", "M", "FN-EQP-05"],
        ["FR-015", "시스템은 신규/개조 설비 도입 시 표준 데이터 매핑 추천 및 일/주/월간 정기 PM 점검 수순서 및 체크리스트를 자동 작성해야 한다.", "M", "FN-ADV-01"],
        ["FR-016", "시스템은 시스템 패치 또는 개조 시 단위/통합 테스트 케이스 및 5단계 시운전 절차서 초안을 자동 수립해야 한다.", "M", "FN-ADV-02"],
        ["FR-017", "시스템은 상위 시스템 다운 시 현장 수동 운전(Manual Override) 전환 절차, 점검 항목, 비상 연락망(BCP 절차서)을 제공해야 한다.", "M", "FN-ADV-03"],
        ["FR-018", "시스템은 개인/팀별 프로젝트 공간을 생성하고 자주 참조하는 매뉴얼/도면을 영구 보관하여 세션 초월 재참조를 지원해야 한다.", "M", "FN-PWS-01"],
        ["FR-019", "시스템은 전체 지식 DB 대신 현재 활성화된 프로젝트 워크스페이스 내 문서로 검색 범위를 한정하여 맞춤형 정밀 검색을 제공해야 한다.", "M", "FN-PWS-02"],
        ["FR-020", "시스템은 필수 인자(설비명/알람코드) 누락 시 에러 표출 대신 선택 칩(Option Chip UI) 형태의 대화형 재확인을 제공해야 한다.", "M", "Concept 2.2.1, 7.1"],
        ["FR-021", "시스템은 지식 검색 기반 응답에 대해 참조된 근거 문서명, 페이지 번호 및 원문 참조 딥링크를 인라인 배너로 표출해야 한다.", "M", "Concept 5.7, 7.2"],
        ["FR-022", "시스템은 검색 근거가 없거나 신뢰도(유사도 0.70) 미달 시 임의 응답 생성을 차단하고 답변 불가 및 수동 확인 가이드를 제공해야 한다.", "M", "Concept 2.2.3, 5.7"]
    ]
    t_fr = doc.add_table(rows=len(fr_data)+1, cols=4)
    style_table(t_fr, [Inches(0.8), Inches(4.2), Inches(0.6), Inches(1.0)], ["ID", "요구사항 명세 내용", "우선순위", "관련 기능/근거"], fr_data, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 5.2 데이터 요구사항 DR
    doc.add_heading("5.2 데이터 요구사항 (Data Requirements - DR)", level=2)
    dr_data = [
        ["DR-001", "시스템은 설비 보전 SOP 및 정비 매뉴얼(PDF)을 청크 임베딩하여 지식 탐색이 가능한 Vector DB로 저장·관리해야 한다.", "M", "지식 관리"],
        ["DR-002", "시스템은 사용자 계정, 역할(User/Admin) 및 프로젝트 워크스페이스 메타데이터를 RDB에 저장·관리해야 한다.", "M", "보안/계정"],
        ["DR-003", "시스템은 사용자 프롬프트 질의 및 피드백 대화 이력을 최소 1년간 안전하게 보관해야 한다.", "M", "이력 관리"],
        ["DR-004", "시스템은 멀티모달 업로드 파일로 PDF, Office(DOCX/PPTX/XLSX), TXT, MD, 이미지를 지원하고 파일당 50MB 이하로 제약해야 한다.", "M", "파일 관리"],
        ["DR-005", "시스템은 개인/팀 프로젝트 워크스페이스의 업로드 파일 메타데이터를 저장·관리해야 한다.", "M", "FN-PWS"],
        ["DR-006", "시스템은 지식 응답별 참조 근거(문서 ID, 페이지 번호, 검색 유사도 점수) 및 답변 불가 처리 여부를 응답 로그에 기록해야 한다.", "M", "품질 검증"]
    ]
    t_dr = doc.add_table(rows=len(dr_data)+1, cols=4)
    style_table(t_dr, [Inches(0.8), Inches(4.2), Inches(0.6), Inches(1.0)], ["ID", "데이터 요구사항 명세", "우선순위", "구분"], dr_data, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 5.3 인터페이스 요구사항 IR
    doc.add_heading("5.3 인터페이스 요구사항 (Interface Requirements - IR)", level=2)
    ir_data = [
        ["IR-001", "에이전트는 현장 설비 PLC/DB와 직접 통신하지 않으며, 모든 외부 데이터 조회는 중계 서버 REST API를 단일 통로로 경유해야 한다.", "M", "연동 안전"],
        ["IR-002", "시스템은 중계 서버 REST API를 통해 MES/WMS로부터 자재 재고, FIFO, 버퍼 체류 시간 및 물류 이송 현황을 조회해야 한다.", "M", "MES/WMS 연동"],
        ["IR-003", "시스템은 중계 서버 REST API를 통해 SCADA/설비 관제 DB로부터 설비 상태, TPH 실적, 알람 로그 및 가동률을 수집해야 한다.", "M", "설비 연동"],
        ["IR-004", "시스템은 SLM의 임의 DB Query(Text-to-SQL) 직접 생성을 엄격히 차단하고, 사전 정의된 중계 서버 API Function Calling만 수행해야 한다.", "M", "DB 과부하 방지"],
        ["IR-005", "시스템은 중계 서버 및 외부 연동 전 구간에 TLS 1.3(HTTPS/WSS) 암호화 통신 프로토콜을 적용해야 한다.", "M", "보안 통신"],
        ["IR-006", "시스템은 현장 태블릿(Android/iOS) 및 관제실 PC(Chrome/Edge) 웹 브라우저를 모두 지원하는 반응형 UI를 제공해야 한다.", "M", "UI 호환성"]
    ]
    t_ir = doc.add_table(rows=len(ir_data)+1, cols=4)
    style_table(t_ir, [Inches(0.8), Inches(4.2), Inches(0.6), Inches(1.0)], ["ID", "인터페이스 요구사항 명세", "우선순위", "구분"], ir_data, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 5.4 성능 및 보안 요구사항 PR & SR
    doc.add_heading("5.4 성능 및 보안 요구사항 (Performance & Security - PR/SR)", level=2)
    pr_sr_data = [
        ["PR-001", "시스템은 단순 질의 및 설비 상태 조회에 대해 2초 이내 응답 피드백을 제공해야 한다.", "M", "응답성"],
        ["PR-002", "시스템은 지식 기반 정비 매뉴얼/SOP 탐색에 대해 4초 이내 응답 피드백을 제공해야 한다.", "M", "지식 검색 성능"],
        ["PR-003", "시스템은 멀티모달 대용량(PDF) 문서 파싱 및 요약에 대해 6초 이내 응답 피드백을 제공해야 한다.", "M", "파싱 성능"],
        ["PR-004", "시스템은 중계 서버 외부 데이터 조회를 1.5초 이내에 완료하고 결과를 전달해야 한다.", "M", "API 성능"],
        ["PR-005", "시스템은 동시 접속 사용자 최소 50 TPS 이상의 처리 성능 및 연중 99.9% 이상의 가용성을 보장해야 한다.", "M", "처리량/SLA"],
        ["SR-001", "시스템은 사용자 인증 및 역할 기반 접근 제어(RBAC: User/Admin)를 제공해야 한다.", "M", "권한 관리"],
        ["SR-002", "시스템은 사내 온프레미스 폐쇄망 환경에서 운영되며, 공장 제어망과는 중계 서버를 통해서만 안전하게 연계되어야 한다.", "M", "망 분리"],
        ["SR-003", "시스템은 비밀번호 복잡도, 변경 주기, 로그인 실패 임계 초과 시 계정 잠금 등 자체 계정 보안 정책을 적용해야 한다.", "M", "계정 보안"],
        ["SR-004", "시스템은 일정 시간 무조작 시 세션을 자동 만료하고 재인증을 요구해야 한다.", "M", "세션 관리"]
    ]
    t_pr_sr = doc.add_table(rows=len(pr_sr_data)+1, cols=4)
    style_table(t_pr_sr, [Inches(0.8), Inches(4.2), Inches(0.6), Inches(1.0)], ["ID", "성능 / 보안 요구사항 명세", "우선순위", "구분"], pr_sr_data, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 5.5 운영 요구사항 OR
    doc.add_heading("5.5 운영 및 유지보수 요구사항 (Operation Requirements - OR)", level=2)
    or_data = [
        ["OR-001", "시스템은 Multi-SLM(Clas, Plan, Proc) 모듈을 컨테이너 기반으로 분리하여 개별 로드 밸런싱 및 확장이 가능해야 한다.", "M", "모듈 확장성"],
        ["OR-002", "시스템은 신규 설비 또는 신규 중계 서버 API 추가 시 SLM 재학습 없이 설정 파일 업데이트만으로 확장이 가능해야 한다.", "M", "유지보수성"],
        ["OR-003", "시스템은 정형 데이터 일 1회 전체 백업 및 15분 증분 백업, 지식 DB 주 1회 전체 백업을 수행해야 한다 (RPO 15분, RTO 30초).", "M", "백업/복구"],
        ["OR-004", "시스템은 설비 개조 및 신규 매뉴얼 도입 시 Vector DB 갱신 관리자 기능을 제공하여 지식 노후화를 방지해야 한다.", "M", "지식 관리"]
    ]
    t_or = doc.add_table(rows=len(or_data)+1, cols=4)
    style_table(t_or, [Inches(0.8), Inches(4.2), Inches(0.6), Inches(1.0)], ["ID", "운영 요구사항 명세", "우선순위", "구분"], or_data, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 6. 외부 시스템 연동 및 UI/UX
    doc.add_heading("6. 외부 시스템 연동 및 UI/UX 컴포넌트 명세", level=1)
    
    doc.add_heading("6.1 외부 연동 대상 및 프로토콜", level=2)
    proto_data = [
        ["에이전트 ↔ 중계 서버", "중계 서버 REST API", "모든 외부 데이터 조회(설비 분석, 가동률, 재고, 알람 등) 단일 통로", "REST API (/api/v1/relay/*)", "단일 안전 관문 (보안 수칙)"],
        ["에이전트 내부", "Vector DB", "설비 보전 SOP, 정비 매뉴얼(PDF), 도면 지식 탐색", "Embedding Vector Search", "전용 지식 탐색"],
        ["중계 서버 ↔ 레거시", "MES / WMS", "자재 입고, FIFO 위반, 공정 버퍼 체류 시간 및 재고 조회", "REST API / DB View", "중계 서버에서 취합 전달"],
        ["중계 서버 ↔ 설비 관제", "SCADA / 설비 DB", "설비 가동 실적, TPH 처리량, 알람 발생/해제 로그 수집", "REST API / DB View", "설비 상태 수집 및 전달"]
    ]
    t_proto = doc.add_table(rows=len(proto_data)+1, cols=5)
    style_table(t_proto, [Inches(1.3), Inches(1.3), Inches(1.8), Inches(1.2), Inches(1.0)], ["연동 구분", "인터페이스 대상", "연동 데이터 및 목적", "연동 프로토콜", "비고"], proto_data, header_bg="1E3A8A")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    doc.add_heading("6.2 주요 UI/UX 컴포넌트", level=2)
    ui_items = [
        ("선택 칩 (Option Chip UI)", "프롬프트 필수 인자 누락 시 에러 대신 보완 선택 항목을 클릭형 버튼 칩으로 즉시 제공"),
        ("멀티모달 파일 업로드 존", "PDF/Office/이미지 첨부 시 파일용량(50MB 이하) 검증 및 파싱 상태 인디케이터 표시"),
        ("프로젝트 워크스페이스 패널", "개인/팀 프로젝트 생성, 담당 설비 매뉴얼 및 도면 등록·관리 패널"),
        ("대화형 표/카드 피드백 존", "SOP 점검 수순, 공정 재고 체류 목록, 가동률 리포트를 Markdown 동적 표 및 요약 카드로 표출"),
        ("원문 출처 인라인 배너", "지식 검색 근거 문서명, 페이지 번호 및 참조 딥링크 컴포넌트")
    ]
    for title, desc in ui_items:
        add_bullet(doc, title, desc)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 7. 요구사항 추적 매트릭스
    doc.add_heading("7. 요구사항 추적 매트릭스 (Traceability Matrix)", level=1)
    add_p(doc, "본 URS에 정의된 기능 요구사항(FR-001 ~ FR-022)과 세부 기능 명세(Features) 및 검증 항목 간의 매핑 관계는 다음과 같다.")
    
    tm_data = [
        ["FR-001", "채팅 인터페이스 & 프롬프트 처리", "FS-UI-01", "FN-SOP-01", "TC-FR-001", "미실시"],
        ["FR-002", "Multi-SLM 파이프라인 엔진", "FS-ENG-01", "Core Engine", "TC-FR-002", "미실시"],
        ["FR-003", "설비 고장 조치 SOP 단계별 가이드", "FS-SOP-01", "FN-SOP-01", "TC-FR-003", "미실시"],
        ["FR-004", "설비 정비 및 부품 매뉴얼 지원", "FS-SOP-02", "FN-SOP-02", "TC-FR-004", "미실시"],
        ["FR-005", "트러블슈팅 조치 이력 탐색", "FS-SOP-03", "FN-SOP-03", "TC-FR-005", "미실시"],
        ["FR-006", "멀티모달 PDF/도면 매뉴얼 융합", "FS-SOP-04", "FN-SOP-04", "TC-FR-006", "미실시"],
        ["FR-007", "물류 이송 설비 이상 진단 및 현황 질의", "FS-EQP-01", "FN-EQP-01", "TC-FR-007", "미실시"],
        ["FR-008", "자동화 이송 시스템 처리 성능 및 병목 진단", "FS-EQP-02", "FN-EQP-02", "TC-FR-008", "미실시"],
        ["FR-009", "적재 및 투입 설비 상태 및 큐 정체 분석", "FS-EQP-03", "FN-EQP-03", "TC-FR-009", "미실시"],
        ["FR-010", "설비 가동률 및 비가동 원인 분석", "FS-EQP-04", "FN-EQP-04", "TC-FR-010", "미실시"],
        ["FR-011", "자재 창고 관리 및 FIFO 추적", "FS-INV-01", "FN-INV-01", "TC-FR-011", "미실시"],
        ["FR-012", "공정 버퍼 체류 시간 및 정체 분석", "FS-INV-02", "FN-INV-02", "TC-FR-012", "미실시"],
        ["FR-013", "공정 간 물류 이송 추적 및 지연 탐지", "FS-INV-03", "FN-INV-03", "TC-FR-013", "미실시"],
        ["FR-014", "설비 알람 통계 및 빈발 장애 분석", "FS-EQP-05", "FN-EQP-05", "TC-FR-014", "미실시"],
        ["FR-015", "표준 데이터 & 정기 PM 체크리스트", "FS-ADV-01", "FN-ADV-01", "TC-FR-015", "미실시"],
        ["FR-016", "테스트 / 5단계 시운전 가이드 지원", "FS-ADV-02", "FN-ADV-02", "TC-FR-016", "미실시"],
        ["FR-017", "비상 대응(BCP) 수동 전환 가이드", "FS-ADV-03", "FN-ADV-03", "TC-FR-017", "미실시"],
        ["FR-018", "개인/팀별 프로젝트 워크스페이스", "FS-PWS-01", "FN-PWS-01", "TC-FR-018", "미실시"],
        ["FR-019", "프로젝트 전용 매뉴얼 맞춤 검색", "FS-PWS-02", "FN-PWS-02", "TC-FR-019", "미실시"],
        ["FR-020", "대화형 선택 칩(Option Chip UI)", "FS-UI-02", "Common UI", "TC-FR-020", "미실시"],
        ["FR-021", "원문 출처 인라인 배너 & 딥링크", "FS-UI-03", "Common UI", "TC-FR-021", "미실시"],
        ["FR-022", "환각 차단 및 답변 불가 안내", "FS-ENG-02", "Grounding", "TC-FR-022", "미실시"]
    ]
    t_tm = doc.add_table(rows=len(tm_data)+1, cols=6)
    style_table(t_tm, [Inches(0.8), Inches(1.8), Inches(0.9), Inches(1.0), Inches(0.9), Inches(0.6)], ["요구사항 ID", "요구사항 명칭", "기능 명세(FS)", "세부 기능(FN)", "테스트 ID", "검증 결과"], tm_data, header_bg="1E3A8A")

    doc.save(output_path)
    print(f"SUCCESS: User-friendly URS DOCX generated -> {output_path}")

if __name__ == '__main__':
    target_docx = r"c:\Users\hanwha\Desktop\PIRELLI\AI AGENT\v0.01\urs_v0_1.docx"
    build_urs_docx(target_docx)
